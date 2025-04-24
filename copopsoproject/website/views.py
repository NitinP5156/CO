import io
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.models import User
from django.contrib.auth.hashers import make_password
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from .forms import ProgramDetailForm, CourseOutcomeForm, CIEDataForm, COPOMappingForm
from .models import ProgramDetail, CourseOutcome, COPOMapping, Student, CIEData
from django.db.models import Avg
from django.http import HttpResponse
from django.http import JsonResponse
from docx import Document
from django.template.loader import render_to_string
from django.contrib.auth.hashers import make_password





def index(request):
    return render(request, 'index.html')

def signup(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password = request.POST['password']

        # Validation checks
        if User.objects.filter(username=username).exists():
            messages.error(request, 'Username already taken.')
        elif User.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered.')
        elif not username or not email or not password:
            messages.error(request, 'All fields are required.')
        else:
            # Create user and redirect to sign in
            User.objects.create(username=username, email=email, password=make_password(password))
            messages.success(request, 'Account created successfully. Please sign in.')
            return redirect('signin')
    return render(request, 'signup.html')

def signin(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')
        else:
            messages.error(request, 'Invalid credentials. Please try again.')
    return render(request, 'signin.html')

@login_required
def home(request):
    # Fetch only the logged-in user's data
    program_details = ProgramDetail.objects.filter(user=request.user)
    return render(request, 'home.html', {'program_details': program_details})

def signout(request):
    logout(request)
    messages.success(request, 'You have been logged out.')
    return redirect('index')

@login_required
def program_details(request):
    if request.method == 'POST':
        form = ProgramDetailForm(request.POST)
        if form.is_valid():
            program = form.save(commit=False)
            program.user = request.user  # Associate program with logged-in user
            program.save()
            return redirect('display')
    else:
        form = ProgramDetailForm()
    return render(request, 'programdetails.html', {'form': form})

@login_required
def edit_program_detail(request, id):
    program = get_object_or_404(ProgramDetail, id=id, user=request.user)  # Restrict access to user-specific data
    if request.method == 'POST':
        form = ProgramDetailForm(request.POST, instance=program)
        if form.is_valid():
            form.save()
            return redirect('display')
    else:
        form = ProgramDetailForm(instance=program)
    return render(request, 'edit_program_detail.html', {'form': form, 'program': program})

@login_required
def delete_program_detail(request, id):
    program = get_object_or_404(ProgramDetail, id=id, user=request.user)  # Restrict access to user-specific data
    if request.method == 'POST':
        program.delete()
        return redirect('display')
    return render(request, 'confirm_delete.html', {'program': program})

@login_required
def course_outcomes(request):
    if request.method == 'POST':
        print(request.POST)  # Debugging
        codes = [
            request.POST.get('code1'),
            request.POST.get('code2'),
            request.POST.get('code3'),
            request.POST.get('code4'),
            request.POST.get('code5'),
        ]
        descriptions = [
            request.POST.get('description1'),
            request.POST.get('description2'),
            request.POST.get('description3'),
            request.POST.get('description4'),
            request.POST.get('description5'),
        ]
        for code, description in zip(codes, descriptions):
            if code and description:
                print(f"Creating CourseOutcome: code={code}, description={description}")
                CourseOutcome.objects.create(code=code, description=description, user=request.user)  # Associate with user
        return redirect('display')
    return render(request, 'course_outcomes.html')

@login_required
def edit_course_outcome(request, id):
    outcome = get_object_or_404(CourseOutcome, id=id, user=request.user)  # Restrict to user-specific data
    if request.method == 'POST':
        form = CourseOutcomeForm(request.POST, instance=outcome)
        if form.is_valid():
            form.save()
            return redirect('display')
    else:
        form = CourseOutcomeForm(instance=outcome)
    return render(request, 'edit_course_outcome.html', {'form': form, 'outcome': outcome})

@login_required
def delete_course_outcome(request, id):
    outcome = get_object_or_404(CourseOutcome, id=id, user=request.user)  # Restrict to user-specific data
    if request.method == 'POST':
        outcome.delete()
        return redirect('display')
    return render(request, 'confirm_delete_course.html', {'outcome': outcome})


@login_required
def copopso_mapping(request):
    if request.method == 'POST':
        rows = ['CO1', 'CO2', 'CO3', 'CO4', 'CO5', 'CO6']
        for row in rows:
            data = {
                'co_number': row,
                'po1': request.POST.get(f'{row}_po1'),
                'po2': request.POST.get(f'{row}_po2'),
                'po3': request.POST.get(f'{row}_po3'),
                'po4': request.POST.get(f'{row}_po4'),
                'po5': request.POST.get(f'{row}_po5'),
                'po6': request.POST.get(f'{row}_po6'),
                'po7': request.POST.get(f'{row}_po7'),
                'po8': request.POST.get(f'{row}_po8'),
                'po9': request.POST.get(f'{row}_po9'),
                'po10': request.POST.get(f'{row}_po10'),
                'po11': request.POST.get(f'{row}_po11'),
                'po12': request.POST.get(f'{row}_po12'),
                'pso1': request.POST.get(f'{row}_pso1'),
                'pso2': request.POST.get(f'{row}_pso2'),
                'pso3': request.POST.get(f'{row}_pso3'),
                'user': request.user,  # Associate with the logged-in user
            }
            # Ensure the entry is unique per user and co_number
            COPOMapping.objects.update_or_create(
                co_number=row, user=request.user, defaults=data
            )
        return redirect('display')

    # Fetch only COPOMappings for the logged-in user
    user_copo_mappings = COPOMapping.objects.filter(user=request.user)
    return render(request, 'copopso.html', {'copo_mappings': user_copo_mappings})



@login_required
def edit_copopso_mapping(request):
    copo_mappings = COPOMapping.objects.filter(user=request.user)  # Filter by user
    
    if request.method == 'POST':
        rows = ['CO1', 'CO2', 'CO3', 'CO4', 'CO5', 'CO6']
        for row in rows:
            data = {
                'co_number': row,
                'po1': request.POST.get(f'{row}_po1'),
                'po2': request.POST.get(f'{row}_po2'),
                'po3': request.POST.get(f'{row}_po3'),
                'po4': request.POST.get(f'{row}_po4'),
                'po5': request.POST.get(f'{row}_po5'),
                'po6': request.POST.get(f'{row}_po6'),
                'po7': request.POST.get(f'{row}_po7'),
                'po8': request.POST.get(f'{row}_po8'),
                'po9': request.POST.get(f'{row}_po9'),
                'po10': request.POST.get(f'{row}_po10'),
                'po11': request.POST.get(f'{row}_po11'),
                'po12': request.POST.get(f'{row}_po12'),
                'pso1': request.POST.get(f'{row}_pso1'),
                'pso2': request.POST.get(f'{row}_pso2'),
                'pso3': request.POST.get(f'{row}_pso3'),
            }
            COPOMapping.objects.update_or_create(
                co_number=row,
                user=request.user,  # Ensure it is user-specific
                defaults=data,
            )
        
        messages.success(request, 'CO-PO-PSO Matrix updated successfully.')
        return redirect('display')

    return render(request, 'edit_copopso_mapping.html', {'copo_mappings': copo_mappings})

@login_required
def delete_copopso_mapping(request, mapping_id):
    mapping = get_object_or_404(COPOMapping, id=mapping_id, user=request.user)  # Restrict to user-specific data
    mapping.delete()
    return redirect('display')


@login_required
def student_details(request):
    if request.method == 'POST':
        student_data = []
        
        # Create students for the logged-in user
        for i in range(1, 6):  
            usn = request.POST.get(f'usn_{i}')
            name = request.POST.get(f'name_{i}')
            if usn and name:
                student_data.append(Student(usn=usn, name=name, user=request.user))  # Associate student with logged-in user

        if student_data:
            Student.objects.bulk_create(student_data, ignore_conflicts=True)
            return redirect('display')  

    rows = range(1, 6) 
    return render(request, 'student.html', {'rows': rows})


@login_required
def edit_student(request, id):
    # Ensure the student belongs to the logged-in user
    student = get_object_or_404(Student, id=id, user=request.user)  # Added user filter
    if request.method == 'POST':
        usn = request.POST.get('usn')
        name = request.POST.get('name')

        if usn and name:
            student.usn = usn
            student.name = name
            student.save()  
            return redirect('display')

        else:
            return render(request, 'edit_student.html', {
                'student': student,
                'error': 'Both USN and Name fields are required.'
            })

    return render(request, 'edit_student.html', {'student': student})


@login_required
def delete_student(request, id):
    # Ensure the student belongs to the logged-in user
    student = get_object_or_404(Student, id=id, user=request.user)  # Added user filter
    if request.method == 'POST':
        student.delete()
        return redirect('display')
    return render(request, 'confirm_delete_student.html', {'student': student})

@login_required
def cie_view(request):
    if request.method == 'POST':
       
        for i in range(1, 5): 
            for j in ['a', 'b', 'c', 'd']: 
                question = f'{i}. {j})'
                first_ia = request.POST.get(f'{i}{j}a')
                second_ia = request.POST.get(f'{i}{j}b')
                third_ia = request.POST.get(f'{i}{j}c')

                # Save to the database with user-specific data
                CIEData.objects.create(
                    question=question,
                    first_ia=first_ia,
                    second_ia=second_ia,
                    third_ia=third_ia,
                    user=request.user  # Associate CIEData with the logged-in user
                )

        return redirect('display')

    return render(request, 'cie.html')


@login_required
def edit_cie_view(request, pk):
    # Ensure the CIEData belongs to the logged-in user
    cie_data = get_object_or_404(CIEData, pk=pk, user=request.user)  # Added user filter

    if request.method == 'POST':
        cie_data.first_ia = request.POST.get('first_ia')
        cie_data.second_ia = request.POST.get('second_ia')
        cie_data.third_ia = request.POST.get('third_ia')
        cie_data.save()

        return redirect('display')

    return render(request, 'edit_cie.html', {'cie_data': cie_data})


@login_required
def delete_cie_view(request, pk):
    # Ensure the CIEData belongs to the logged-in user
    cie_data = get_object_or_404(CIEData, pk=pk, user=request.user)  # Added user filter

    if request.method == 'POST':
        cie_data.delete()
        return redirect('display')

    return render(request, 'delete_cie.html', {'cie_data': cie_data})



@login_required
def display(request):
    # Filter data by the logged-in user
    program_details = ProgramDetail.objects.filter(user=request.user)
    course_outcomes = CourseOutcome.objects.filter(user=request.user)
    students = Student.objects.filter(user=request.user)
    copo_mappings = COPOMapping.objects.filter(user=request.user)
    cie_data = CIEData.objects.filter(user=request.user)
    
    # Calculate averages for user-specific COPOMappings
    averages = {
        'po1': copo_mappings.aggregate(Avg('po1'))['po1__avg'] or 0,
        'po2': copo_mappings.aggregate(Avg('po2'))['po2__avg'] or 0,
        'po3': copo_mappings.aggregate(Avg('po3'))['po3__avg'] or 0,
        'po4': copo_mappings.aggregate(Avg('po4'))['po4__avg'] or 0,
        'po5': copo_mappings.aggregate(Avg('po5'))['po5__avg'] or 0,
        'po6': copo_mappings.aggregate(Avg('po6'))['po6__avg'] or 0,
        'po7': copo_mappings.aggregate(Avg('po7'))['po7__avg'] or 0,
        'po8': copo_mappings.aggregate(Avg('po8'))['po8__avg'] or 0,
        'po9': copo_mappings.aggregate(Avg('po9'))['po9__avg'] or 0,
        'po10': copo_mappings.aggregate(Avg('po10'))['po10__avg'] or 0,
        'po11': copo_mappings.aggregate(Avg('po11'))['po11__avg'] or 0,
        'po12': copo_mappings.aggregate(Avg('po12'))['po12__avg'] or 0,
        'pso1': copo_mappings.aggregate(Avg('pso1'))['pso1__avg'] or 0,
        'pso2': copo_mappings.aggregate(Avg('pso2'))['pso2__avg'] or 0,
        'pso3': copo_mappings.aggregate(Avg('pso3'))['pso3__avg'] or 0,
    }

    # Handle POST requests for editing and deleting students
    if request.method == 'POST':
        if 'edit_student' in request.POST:
            student_id = request.POST.get('student_id')
            student = get_object_or_404(Student, id=student_id, user=request.user)
            student.usn = request.POST.get('usn', student.usn)
            student.name = request.POST.get('name', student.name)
            student.save()
        elif 'delete_student' in request.POST:
            student_id = request.POST.get('student_id')
            student = get_object_or_404(Student, id=student_id, user=request.user)
            student.delete()

        return redirect('display')

    # Render the template with user-specific data
    return render(request, 'display.html', {
        'program_details': program_details,
        'course_outcomes': course_outcomes,
        'copo_mappings': copo_mappings,
        'averages': averages,
        'students': students,
        'cie_data': cie_data,
    })






@login_required
def generate_docx(request):
    # Create a Document object
    doc = Document()

    # Title of the document
    doc.add_heading('All The Details', 0)

    # Get the currently logged-in user
    user = request.user

    # Add Course Outcomes to the document (Auto-fetch from database for the logged-in user)
    doc.add_heading('Course Outcomes', level=1)
    course_outcomes = CourseOutcome.objects.filter(user=user)  # Filter by logged-in user
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Code'
    hdr_cells[1].text = 'Description'
    
    for outcome in course_outcomes:
        row_cells = table.add_row().cells
        row_cells[0].text = str(outcome.code)  # Ensure code is a string
        row_cells[1].text = str(outcome.description)  # Ensure description is a string

    # Add Student Details (Auto-fetch from database for the logged-in user)
    doc.add_heading('Student Details', level=1)
    students = Student.objects.filter(user=user)  # Filter by logged-in user
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sl.No'
    hdr_cells[1].text = 'USN'
    hdr_cells[2].text = 'Name'
    
    for idx, student in enumerate(students, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)  # Ensure index is a string
        row_cells[1].text = str(student.usn)  # Ensure USN is a string
        row_cells[2].text = str(student.name)  # Ensure name is a string

    # Add CIE Data (Auto-fetch from database for the logged-in user)
    doc.add_heading('CIE Data', level=1)
    cie_data = CIEData.objects.filter(user=user)  # Filter by logged-in user
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Q. No.'
    hdr_cells[1].text = '1st IA'
    hdr_cells[2].text = '2nd IA'
    hdr_cells[3].text = '3rd IA'
    
    for data in cie_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(data.question)  # Ensure question is a string
        row_cells[1].text = str(data.first_ia)  # Ensure first IA is a string
        row_cells[2].text = str(data.second_ia)  # Ensure second IA is a string
        row_cells[3].text = str(data.third_ia)  # Ensure third IA is a string

    # Add CO-PO-PSO Mapping (Auto-fetch from database for the logged-in user)
    doc.add_heading('CO-PO-PSO Mapping', level=1)
    copo_mappings = COPOMapping.objects.filter(user=user)  # Filter by logged-in user
    table = doc.add_table(rows=1, cols=16)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'CO/PO & PSO'
    for i in range(1, 16):
        hdr_cells[i].text = f'PO{i}' if i <= 12 else f'PSO{i-12}'
    
    for mapping in copo_mappings:
        row_cells = table.add_row().cells
        row_cells[0].text = str(mapping.co_number)  # Ensure co_number is a string
        for i in range(1, 16):
            row_cells[i].text = str(getattr(mapping, f'po{i}', '') or getattr(mapping, f'pso{i-12}', ''))  # Ensure each PO/PSO is a string

    # Add Program Details (Auto-fetch from database for the logged-in user)
    doc.add_heading('Program Details', level=1)
    program_details = ProgramDetail.objects.filter(user=user)  # Filter by logged-in user
    table = doc.add_table(rows=1, cols=10)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'NBA Code'
    hdr_cells[1].text = 'Dept'
    hdr_cells[2].text = 'Course Name'
    hdr_cells[3].text = 'Academic Year'
    hdr_cells[4].text = 'No of Students'
    hdr_cells[5].text = 'Faculty Name'
    hdr_cells[6].text = 'Program Title'
    hdr_cells[7].text = 'Semester & Section'
    hdr_cells[8].text = 'Max SEE Marks'
    hdr_cells[9].text = 'HoD Name'
    
    for program in program_details:
        row_cells = table.add_row().cells
        row_cells[0].text = str(program.nba_code)  # Ensure nba_code is a string
        row_cells[1].text = str(program.department)  # Ensure department is a string
        row_cells[2].text = str(program.course_name)  # Ensure course_name is a string
        row_cells[3].text = str(program.academic_year)  # Ensure academic_year is a string
        row_cells[4].text = str(program.num_students)  # Ensure num_students is a string
        row_cells[5].text = str(program.faculty_name)  # Ensure faculty_name is a string
        row_cells[6].text = str(program.program_title)  # Ensure program_title is a string
        row_cells[7].text = str(program.semester_section)  # Ensure semester_section is a string
        row_cells[8].text = str(program.max_see_marks)  # Ensure max_see_marks is a string
        row_cells[9].text = str(program.hod_name)  # Ensure hod_name is a string

    # Create the HTTP response to download the file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=program_details.docx'
    doc.save(response)
    return response
